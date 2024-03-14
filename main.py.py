import imaplib
import email
import os 
from docx import Document

user = 'madhuprincess057@gmail.com'
password = 'btxfdlbcivgztunt'
imap_url = 'imap.gmail.com'
attachment_dir = r'path/to/your/attachment/dir'

def auth(user, password, imap_url):
    con = imaplib.IMAP4_SSL(imap_url)
    con.login(user, password)
    return con

def get_body(msg):
    if msg.is_multipart():
        return get_body(msg.get_payload(0))
    else:
        return msg.get_payload(None, True)

def get_attachments(msg, con):
    for part in msg.walk():
        if part.get_content_maintype() == 'multipart':
            continue
        if part.get('Content-Disposition') is None:
            continue
        fileName = part.get_filename()

        if bool(fileName):
            filePath = os.path.join(attachment_dir, fileName)
            with open(filePath, 'wb') as f:
                f.write(part.get_payload(decode=True))

def search(key, value, con):
    result, data = con.search(None, key, '"{}"'.format(value))
    return data

def get_emails(result_bytes, con):
    msgs = []
    for num in result_bytes[0].split():
        typ, data = con.fetch(num, '(RFC822)')
        msgs.append(data)
    return msgs

if not os.path.exists(attachment_dir):
    os.makedirs(attachment_dir)

con = auth(user, password, imap_url)
con.select('INBOX')

result, data = con.search(None, 'ALL')
emails = get_emails(data, con)

for email_data in emails:
    raw = email.message_from_bytes(email_data[1])
    get_attachments(raw, con)

# Create a new Word document
doc = Document()
doc.add_heading('Attachment Directory', level=1)
doc.add_paragraph(f'The attachment directory is located at: {attachment_dir}')
doc.save('attachment_directory.docx')

print('Word document created successfully.')